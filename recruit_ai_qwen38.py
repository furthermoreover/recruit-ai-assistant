from fastapi import FastAPI,Body,HTTPException,Depends,Header,UploadFile,File
from fastapi.responses import RedirectResponse
import pandas as pd
from typing import Optional,Dict
import json
import os
import requests
import hashlib, secrets, time, threading, shutil
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

base_url = os.getenv("LLM_BASE_URL")
api_key = os.getenv("LLM_API_KEY")
llm_model = os.getenv("LLM_MODEL")

if not all([base_url, api_key, llm_model]):
    raise RuntimeError(".env请配置 LLM_BASE_URL / LLM_API_KEY / LLM_MODEL")

client = OpenAI(
    api_key=api_key,
    base_url=base_url
)

# ===== 可切换的全局模型（前端可选） =====
# 启动时用 .env 的模型作为默认；可在前端通过 /model 接口动态切换，无需重启。
CURRENT_MODEL = llm_model

def _fetch_available_models():
    """尝试从 /models 接口拉取账号可用模型，失败时回退内置候选列表。"""
    fallback = ["qwen3.7-max", "qwen3.7-plus", "qwen3.6-flash", "glm-5.2",
                "deepseek-v4-pro", "deepseek-v4-flash-0731",
                "qwen3.8-max", "qwen3.8-flash"]
    try:
        r = requests.get(base_url.rstrip("/") + "/models",
                         headers={"Authorization": "Bearer " + api_key}, timeout=10)
        if r.status_code == 200:
            ids = [m["id"] for m in r.json().get("data", [])]
            if ids:
                return ids
    except Exception:
        pass
    return fallback

AVAILABLE_MODELS = _fetch_available_models()

# ============================================================
# 用户系统：注册 / 登录 / token 鉴权 / 数据按用户隔离
# 每个用户一个独立台账文件：data/<username>/recruit_record.xlsx
# ============================================================
# 数据目录：默认与源码同目录的 data/；exe 打包时由 launcher 通过环境变量 RECRUIT_DATA_DIR 指向 exe 同目录
DATA_DIR = os.environ.get("RECRUIT_DATA_DIR") or os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
USER_FILE = os.path.join(DATA_DIR, "users.json")
_tokens = {}                      # token -> username（内存态，重启需重新登录）
_tokens_lock = threading.Lock()

def _load_users():
    if os.path.exists(USER_FILE):
        try:
            with open(USER_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def _save_users(users):
    os.makedirs(DATA_DIR, exist_ok=True)
    with open(USER_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=2)

def _hash_pwd(pwd, salt):
    return hashlib.pbkdf2_hmac("sha256", pwd.encode("utf-8"), salt.encode("utf-8"), 100000).hex()

def _user_dir(username):
    d = os.path.join(DATA_DIR, username)
    os.makedirs(d, exist_ok=True)
    return d

def get_current_user(authorization: str = Header(default="")):
    """FastAPI 依赖：从 Authorization: Bearer <token> 解析当前登录用户。"""
    token = ""
    if authorization.startswith("Bearer "):
        token = authorization[7:].strip()
    with _tokens_lock:
        user = _tokens.get(token) if token else None
    if not user:
        raise HTTPException(status_code=401, detail="未登录或登录已过期")
    return user

def _init_admin():
    """确保 admin 用户存在（初始密码 123456），并把根目录旧台账迁移到 admin。"""
    users = _load_users()
    if "admin" not in users:
        salt = secrets.token_hex(8)
        users["admin"] = {"salt": salt, "hash": _hash_pwd("123456", salt),
                          "created": time.strftime("%Y-%m-%d %H:%M:%S")}
        _save_users(users)
    admin_file = os.path.join(_user_dir("admin"), "recruit_record.xlsx")
    if not os.path.exists(admin_file) and os.path.exists("recruit_record.xlsx"):
        try:
            shutil.copy("recruit_record.xlsx", admin_file)
        except Exception:
            pass

_init_admin()

app = FastAPI(title="招聘助理AI｜OpenAI兼容接口")

@app.get("/", include_in_schema=False)
async def root():
    return RedirectResponse(url="/docs")

@app.get("/model", include_in_schema=False)
def get_model():
    return {"current": CURRENT_MODEL, "models": AVAILABLE_MODELS}

@app.post("/model", include_in_schema=False)
def set_model(body: dict = Body(...)):
    global CURRENT_MODEL
    m = (body.get("model") or "").strip()
    if m in AVAILABLE_MODELS:
        CURRENT_MODEL = m
        return {"ok": True, "current": CURRENT_MODEL}
    return {"ok": False, "msg": f"未知模型: {m}"}


# ========== 注册 / 登录 / 登出 ==========
@app.post("/register", include_in_schema=False)
def register(body: dict = Body(...)):
    username = (body.get("username") or "").strip()
    password = str(body.get("password") or "")
    if not username or not password:
        return {"ok": False, "msg": "用户名和密码不能为空"}
    if len(username) < 2 or len(username) > 20:
        return {"ok": False, "msg": "用户名需 2-20 个字符"}
    if len(password) < 6:
        return {"ok": False, "msg": "密码至少 6 位"}
    users = _load_users()
    if username in users:
        return {"ok": False, "msg": "用户名已存在"}
    salt = secrets.token_hex(8)
    users[username] = {"salt": salt, "hash": _hash_pwd(password, salt),
                       "created": time.strftime("%Y-%m-%d %H:%M:%S")}
    _save_users(users)
    return {"ok": True, "msg": "注册成功，请登录"}


@app.post("/login", include_in_schema=False)
def login(body: dict = Body(...)):
    username = (body.get("username") or "").strip()
    password = str(body.get("password") or "")
    users = _load_users()
    u = users.get(username)
    if not u or _hash_pwd(password, u["salt"]) != u["hash"]:
        return {"ok": False, "msg": "用户名或密码错误"}
    token = secrets.token_hex(16)
    with _tokens_lock:
        _tokens[token] = username
    return {"ok": True, "token": token, "username": username}


@app.post("/logout", include_in_schema=False)
def logout(body: dict = Body(...)):
    token = (body.get("token") or "").strip()
    with _tokens_lock:
        _tokens.pop(token, None)
    return {"ok": True}


@app.get("/me", include_in_schema=False)
def me(user: str = Depends(get_current_user)):
    return {"username": user}


# ========== 文件上传解析（PDF / Word / TXT → 文本） ==========
def extract_file_text(content: bytes, filename: str) -> str:
    """按扩展名提取 PDF / Word(.docx) / TXT 文本。"""
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        from io import BytesIO
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if lower.endswith(".docx"):
        from io import BytesIO
        from docx import Document
        doc = Document(BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if lower.endswith(".txt") or lower.endswith(".md"):
        return content.decode("utf-8", errors="replace")
    if lower.endswith(".doc"):
        raise HTTPException(status_code=400, detail="暂不支持老版 .doc 格式，请另存为 .docx 或 PDF")
    raise HTTPException(status_code=400, detail="暂不支持该文件格式，请上传 PDF / Word(.docx) / TXT")


@app.post("/upload_file", include_in_schema=False)
async def upload_file(file: UploadFile = File(...), user: str = Depends(get_current_user)):
    """上传简历文档（PDF/Word/TXT），提取文本返回，供简历录入与 AI 助理使用。"""
    content = await file.read()
    filename = file.filename or "file"
    text = extract_file_text(content, filename)
    if len(text) > 30000:
        text = text[:30000] + "\n...[内容过长已截断]"
    return {"filename": filename, "text": text, "chars": len(text)}

EXCEL_FILE = "recruit_record.xlsx"

RESUME_SCHEMA = {
    "type":"object",
    "properties":{
        "candidate_name":{"type":"string","description":"候选人姓名，识别不到填未获取"},
        "phone":{"type":"string","description":"手机号，识别不到填未获取"},
        "expect_salary":{"type":"string","description":"期望薪资，识别不到填未获取"},
        "available_time":{"type":"string","description":"到岗时间，识别不到填未获取"}
    },
    "required":["candidate_name","phone","expect_salary","available_time"]
}

def load_excel(username="admin"):
    file_path = os.path.join(_user_dir(username), "recruit_record.xlsx")
    if os.path.exists(file_path):
        # 全部按字符串读取，避免手机号/期望薪资被读成 float(如 11125156481.0) 影响匹配
        df = pd.read_excel(file_path, dtype=str)
        # 把读出来的字符串"nan"/NaN替换为空
        df = df.replace("nan", "").replace(float("nan"), "")
    else:
        # 新建空表逻辑（列名须与 resume_parse_add 写入的 new_row 一致）
        df = pd.DataFrame(columns=[
            "岗位名称","候选人姓名","手机号","招聘渠道","期望薪资","到岗时间",
            "沟通状态","面试时间","面试形式","面试结果","备注"
        ])
    return df


def save_excel(df, username="admin"):
    file_path = os.path.join(_user_dir(username), "recruit_record.xlsx")
    df.to_excel(file_path, index=False, engine="openpyxl")


def _norm(s):
    """规范化字符串：去首尾空白；去掉浮点转换带来的 .0 后缀（如 11125156481.0）"""
    s = "" if s is None else str(s).strip()
    if s.endswith(".0") and s[:-2].isdigit():
        s = s[:-2]
    return s


def _dup_mask(df, name, phone):
    """返回与 (name, phone) 冲突的行的布尔 Series：姓名相同 或 手机号相同(且非空) 即视为同一候选人"""
    if df is None or len(df) == 0:
        return pd.Series(dtype=bool)
    name = _norm(name)
    phone = _norm(phone)
    names = df["候选人姓名"].map(_norm)
    phones = df["手机号"].map(_norm)
    mask = pd.Series(False, index=df.index)
    if name and name != "未获取":
        mask = mask | (names == name)
    if phone and phone != "未获取":
        mask = mask | (phones == phone)
    return mask

# ==========接口1：生成招聘沟通话术 ==========
@app.post("/gen_talk")
async def gen_talk(
    post_name:str=Body(...,description="招聘岗位名称"),
    scene:str=Body(...,description="场景：打招呼 /邀约面试 /候选人爽约 /婉拒候选人 /复试通知"),
    candidate_info:Optional[Dict]=Body(default=None,description="候选人简单信息字典"),
    user: str = Depends(get_current_user)
):
    sys_prompt ="""
你是HR招聘助理。
只根据用户传入的【使用场景】，**仅仅输出这一个场景对应的1条BOSS直聘消息**。
禁止输出多个场景、禁止输出多条。输出简短、适合复制粘贴，纯文本，不要markdown，不要多余解释。
"""
    user_text = f"""
岗位：{post_name}
使用场景：{scene}
候选人信息：{candidate_info if candidate_info else "无"}
"""
    try:
        resp = client.chat.completions.create(
            model=CURRENT_MODEL,
            messages=[
                {"role":"system","content":sys_prompt},
                {"role":"user","content":user_text}
            ],
            max_tokens=512
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"大模型调用异常:{str(e)}")

    content = resp.choices[0].message.content
    return {"talk_content":content.strip()}


# ==========接口2：简历文本解析入库（支持重复检测/更新/强制新增） ==========
@app.post("/resume_parse_add")
async def resume_parse_add(
    post_name:str=Body(...),
    channel:str=Body(...),
    resume_text:str=Body(...,description="粘贴简历文本"),
    mode:str=Body("auto",description="auto=自动查重后新增；update=更新已有候选人；new=强制新增"),
    target_name:Optional[str]=Body(None,description="mode=update 时指定要更新的候选人姓名"),
    target_phone:Optional[str]=Body(None,description="mode=update 时指定要更新的候选人手机号(可选，用于精确匹配)"),
    target_key:Optional[int]=Body(None,description="mode=update 时指定要更新的候选人行索引key(优先于 target_name)"),
    user: str = Depends(get_current_user)
):
    sys_prompt = """
你是简历信息抽取助手。只输出严格JSON对象，不要输出任何其他文字、解释、markdown。
字段一共4个：
candidate_name：候选人姓名
phone：联系电话
expect_salary：期望薪资
available_time：到岗时间
识别不到的字段，值填写字符串“未获取”。
输出示例：{"candidate_name":"张三","phone":"13800138000","expect_salary":"8‑10k","available_time":"一周内"}
"""
    try:
        resp = client.chat.completions.create(
            model=CURRENT_MODEL,
            messages=[
                {"role":"system","content":sys_prompt},
                {"role":"user","content":f"简历内容：{resume_text}"}
            ],
            max_tokens=800
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"简历解析调用异常:{str(e)}")

    json_str = resp.choices[0].message.content
    try:
        parse_data = json.loads(json_str)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"JSON解析失败:{str(e)},raw:{json_str}")

    df = load_excel(user)
    cand_name = _norm(parse_data.get("candidate_name", ""))
    cand_phone = _norm(parse_data.get("phone", ""))

    # ---------- mode=update：更新已有候选人（前端下拉选中后调用） ----------
    if mode == "update":
        if target_key is not None:
            # 按行索引key精确定位（避免"未获取"/重名误匹配）
            if target_key < 0 or target_key >= len(df):
                raise HTTPException(status_code=404, detail="目标候选人记录不存在或已被删除")
            mask = pd.Series(False, index=df.index)
            mask.iloc[target_key] = True
            match_name = _norm(df.iloc[target_key]["候选人姓名"])
        else:
            match_name = _norm(target_name or cand_name)
            match_phone = _norm(target_phone or cand_phone)
            if not match_name:
                raise HTTPException(status_code=400, detail="update 模式需要提供候选人姓名")
            names = df["候选人姓名"].map(_norm)
            mask = (names == match_name)
            # 仅当未显式指定 target_name(即按本次解析结果匹配)时，才用手机号辅助匹配
            if not target_name:
                phones = df["手机号"].map(_norm)
                if match_phone and match_phone != "未获取":
                    mask = mask | (phones == match_phone)
        if mask.sum() == 0:
            raise HTTPException(status_code=404, detail=f"未找到候选人【{match_name}】的记录，无法更新")

        # 合并字段：新值有效(非空/非"未获取")时覆盖旧值；新值无效则保留原值
        def _merge(col, new_val):
            new_val = _norm(new_val)
            if not new_val or new_val == "未获取":
                return
            df.loc[mask, col] = new_val

        _merge("候选人姓名", cand_name)
        _merge("手机号", cand_phone)
        _merge("期望薪资", parse_data.get("expect_salary", ""))
        _merge("到岗时间", parse_data.get("available_time", ""))
        _merge("岗位名称", post_name)
        _merge("招聘渠道", channel)
        # 追加更新备注
        old_note = _norm(df.loc[mask, "备注"].iloc[0])
        df.loc[mask, "备注"] = f"{old_note}\n【更新】{resume_text[:200]}"
        save_excel(df, user)
        return {"msg": f"已更新候选人【{match_name}】的信息", "updated": True, "parsed": parse_data}

    # ---------- mode=auto：自动查重，重复则提示不写入 ----------
    dup_mask = _dup_mask(df, cand_name, cand_phone)
    if mode == "auto" and dup_mask.sum() > 0:
        dup_rows = []
        for i in df.index[dup_mask]:
            dup_rows.append({
                "name": _norm(df.loc[i, "候选人姓名"]),
                "phone": _norm(df.loc[i, "手机号"]),
                "expect_salary": _norm(df.loc[i, "期望薪资"]),
                "available_time": _norm(df.loc[i, "到岗时间"]),
                "status": _norm(df.loc[i, "沟通状态"]),
            })
        return {"msg": "检测到该候选人已存在，未写入台账", "duplicate": True, "exists": dup_rows}

    # ---------- 新增一行 ----------
    new_row = {
        "岗位名称":post_name,
        "候选人姓名":cand_name,
        "手机号":cand_phone,
        "招聘渠道":channel,
        "期望薪资":_norm(parse_data.get("expect_salary", "")),
        "到岗时间":_norm(parse_data.get("available_time", "")),
        "沟通状态":"待沟通",
        "面试时间":"",
        "面试形式":"",
        "面试结果":"",
        "备注":resume_text[:300]
    }
    df.loc[len(df)] = new_row
    save_excel(df, user)
    return {"msg":"简历解析存入台账成功","parsed":parse_data}

# ==========接口3：更新候选人状态 ==========
@app.post("/update_status")
async def update_status(
    name: str=Body(...,description="候选人姓名"),
    status: str=Body("",description="沟通状态"),
    interview_time: str=Body("",description="面试时间"),
    interview_type: str=Body("",description="面试形式"),
    result: str=Body("",description="面试结果"),
    note: str=Body("",description="备注信息"),
    key: Optional[int]=Body(None,description="候选人记录的行索引key(优先于name精确定位)"),
    user: str = Depends(get_current_user)
):
    df = load_excel(user)

    # 关键修复：强制转为字符串，避免float64不能存字符串报错
    for col in ["沟通状态","面试时间","面试形式","面试结果","备注"]:
        df[col] = df[col].astype(str).replace("nan","")

    if key is not None:
        # 按行索引key精确定位
        if key < 0 or key >= len(df):
            raise HTTPException(status_code=404,detail="该候选人记录不存在或已被删除")
        mask = pd.Series(False, index=df.index)
        mask.iloc[key] = True
    else:
        mask = df["候选人姓名"] == name
        if mask.sum() == 0:
            raise HTTPException(status_code=404,detail=f"没有找到姓名为【{name}】的候选人记录")

    df.loc[mask,"沟通状态"] = status
    df.loc[mask,"面试时间"] = interview_time
    df.loc[mask,"面试形式"] = interview_type
    df.loc[mask,"面试结果"] = result

    if note.strip()!="":
        old_note = df.loc[mask,"备注"].iloc[0]
        df.loc[mask,"备注"] = f"{old_note}\n【更新备注】{note}"

    save_excel(df, user)
    return {"msg":"状态更新成功","candidate_name":name}


# ==========接口4：读取全部台账 ==========
@app.get("/get_all_record")
async def get_all_record(user: str = Depends(get_current_user)):
    df = load_excel(user)
    return df.to_dict(orient="records")


# ==========接口：获取全部已录入候选人（含"未获取"，带唯一key供下拉选择/删除） ==========
@app.get("/get_candidate_names")
async def get_candidate_names(user: str = Depends(get_current_user)):
    df = load_excel(user)
    if len(df) == 0:
        return []
    result = []
    for i in range(len(df)):
        result.append({
            "key": i,
            "name": _norm(df.loc[i, "候选人姓名"]),
            "phone": _norm(df.loc[i, "手机号"]),
            "post": _norm(df.loc[i, "岗位名称"]),
        })
    return result


# ==========接口：删除候选人记录（按行索引key精确定位） ==========
@app.post("/delete_candidate")
async def delete_candidate(key: int = Body(..., embed=True, description="候选人记录的行索引key"),
                           user: str = Depends(get_current_user)):
    df = load_excel(user)
    if len(df) == 0 or key < 0 or key >= len(df):
        raise HTTPException(status_code=404, detail="该记录不存在或已被删除")
    row = df.iloc[key]
    deleted = {
        "name": _norm(row["候选人姓名"]),
        "phone": _norm(row["手机号"]),
        "post": _norm(row["岗位名称"]),
    }
    df = df.drop(index=df.index[key]).reset_index(drop=True)
    save_excel(df, user)
    return {"msg": f"已删除候选人【{deleted['name']}】（{deleted['phone']}）的记录", "deleted": deleted}


# ==========接口：获取已录入岗位列表（去重，供前端岗位下拉选择） ==========
@app.get("/get_job_list")
async def get_job_list(user: str = Depends(get_current_user)):
    df = load_excel(user)
    if len(df) == 0:
        return []
    jobs = []
    seen = set()
    for i in range(len(df)):
        j = _norm(df.loc[i, "岗位名称"])
        if not j or j in seen:
            continue
        seen.add(j)
        jobs.append(j)
    return jobs


# ==========接口5：统计（供前端"招聘统计"Tab调用） ==========
@app.get("/stat_count")
async def stat_count(user: str = Depends(get_current_user)):
    df = load_excel(user)
    total = len(df)
    by_status = df["沟通状态"].replace("", "未设置").value_counts().to_dict()
    by_result = df["面试结果"].replace("", "待评价").value_counts().to_dict()
    return {
        "total": total,
        "by_status": by_status,
        "by_result": by_result
    }

# ============================================================
# ========== Agent 对话式指挥（LLM 工具调用） ==========
# ============================================================

AGENT_TOOLS = [
    {"type": "function", "function": {"name": "list_candidates", "description": "列出台账中所有已录入候选人（含未获取），含key/姓名/手机号/岗位", "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {"name": "get_all_records", "description": "获取完整候选人台账明细（所有字段）", "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {"name": "get_jobs", "description": "获取已录入岗位列表（去重）", "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {"name": "get_stats", "description": "获取招聘统计：总数、按沟通状态、按面试结果", "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {"name": "gen_talk", "description": "为指定岗位和场景生成一条BOSS直聘沟通话术", "parameters": {"type": "object", "properties": {"post_name": {"type": "string", "description": "岗位名称"}, "scene": {"type": "string", "description": "场景，如：初次打招呼/邀约面试/跟进回复/婉拒候选人"}, "candidate_info": {"type": "object", "description": "候选人信息（可选）"}}, "required": ["post_name", "scene"]}}},
    {"type": "function", "function": {"name": "parse_resume_add", "description": "粘贴简历文本，抽取姓名/手机号/薪资/到岗时间写入台账；mode=auto查重/update更新/new新增", "parameters": {"type": "object", "properties": {"post_name": {"type": "string", "description": "岗位名称"}, "channel": {"type": "string", "description": "简历来源"}, "resume_text": {"type": "string", "description": "简历文本"}, "mode": {"type": "string", "enum": ["auto", "update", "new"], "description": "auto/update/new"}, "target_key": {"type": "integer", "description": "update时指定候选人行索引key"}}, "required": ["post_name", "channel", "resume_text"]}}},
    {"type": "function", "function": {"name": "update_candidate_status", "description": "更新候选人状态：沟通状态/面试时间/面试形式/面试结果/备注", "parameters": {"type": "object", "properties": {"name": {"type": "string", "description": "候选人姓名（与key二选一）"}, "key": {"type": "integer", "description": "候选人行索引key（优先于name）"}, "status": {"type": "string", "description": "沟通状态"}, "interview_time": {"type": "string", "description": "面试时间"}, "interview_type": {"type": "string", "description": "面试形式"}, "result": {"type": "string", "description": "面试结果"}, "note": {"type": "string", "description": "备注"}}}}},
    {"type": "function", "function": {"name": "delete_candidate", "description": "按行索引key删除候选人记录，删除前先确认", "parameters": {"type": "object", "properties": {"key": {"type": "integer", "description": "要删除的记录行索引key"}}, "required": ["key"]}}}
]


def _agent_tool(name, args, user="admin"):
    """执行 Agent 工具，返回可 JSON 序列化的结果 dict（按 user 隔离台账）"""
    # 1. 列出候选人
    if name == "list_candidates":
        df = load_excel(user)
        result = []
        for i in range(len(df)):
            result.append({
                "key": i,
                "name": _norm(df.loc[i, "候选人姓名"]),
                "phone": _norm(df.loc[i, "手机号"]),
                "post": _norm(df.loc[i, "岗位名称"]),
            })
        return {"candidates": result}

    # 2. 完整台账
    if name == "get_all_records":
        df = load_excel(user)
        return {"records": df.to_dict(orient="records")}

    # 3. 岗位列表
    if name == "get_jobs":
        df = load_excel(user)
        jobs, seen = [], set()
        for i in range(len(df)):
            j = _norm(df.loc[i, "岗位名称"])
            if j and j not in seen:
                seen.add(j)
                jobs.append(j)
        return {"jobs": jobs}

    # 4. 统计
    if name == "get_stats":
        df = load_excel(user)
        total = len(df)
        by_status = df["沟通状态"].replace("", "未设置").value_counts().to_dict()
        by_result = df["面试结果"].replace("", "待评价").value_counts().to_dict()
        return {"total": total, "by_status": by_status, "by_result": by_result}

    # 5. 生成话术
    if name == "gen_talk":
        post_name = str(args.get("post_name", ""))
        scene = str(args.get("scene", ""))
        candidate_info = args.get("candidate_info")
        sys_prompt = """
你是HR招聘助理。
只根据用户传入的【使用场景】，**仅仅输出这一个场景对应的1条BOSS直聘消息**。
禁止输出多个场景、禁止输出多条。输出简短、适合复制粘贴，纯文本，不要markdown，不要多余解释。
"""
        user_text = f"岗位：{post_name}\n使用场景：{scene}\n候选人信息：{candidate_info if candidate_info else '无'}"
        resp = client.chat.completions.create(
            model=CURRENT_MODEL,
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": user_text}
            ],
            max_tokens=512
        )
        return {"talk_content": resp.choices[0].message.content.strip()}

    # 6. 简历解析入库
    if name == "parse_resume_add":
        post_name = str(args.get("post_name", ""))
        channel = str(args.get("channel", ""))
        resume_text = str(args.get("resume_text", ""))
        mode = str(args.get("mode", "auto"))
        target_key = args.get("target_key")
        target_name = args.get("target_name")
        target_phone = args.get("target_phone")
        if not resume_text.strip():
            return {"msg": "简历文本为空，无法解析", "parsed": None}
        sys_prompt = """
你是简历信息抽取助手。只输出严格JSON对象，不要输出任何其他文字、解释、markdown。
字段一共4个：
candidate_name：候选人姓名
phone：联系电话
expect_salary：期望薪资
available_time：到岗时间
识别不到的字段，值填写字符串“未获取”。
输出示例：{"candidate_name":"张三","phone":"13800138000","expect_salary":"8-10k","available_time":"一周内"}
"""
        resp = client.chat.completions.create(
            model=CURRENT_MODEL,
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": f"简历内容：{resume_text}"}
            ],
            max_tokens=800
        )
        try:
            parse_data = json.loads(resp.choices[0].message.content)
        except json.JSONDecodeError:
            return {"msg": "简历解析失败：模型输出不是合法JSON", "parsed": None}
        df = load_excel(user)
        cand_name = _norm(parse_data.get("candidate_name", ""))
        cand_phone = _norm(parse_data.get("phone", ""))
        if mode == "update":
            if target_key is not None:
                if target_key < 0 or target_key >= len(df):
                    return {"msg": f"候选人记录(key={target_key})不存在或已被删除"}
                mask = pd.Series(False, index=df.index)
                mask.iloc[target_key] = True
                match_name = _norm(df.iloc[target_key]["候选人姓名"])
            else:
                match_name = _norm(target_name or cand_name)
                match_phone = _norm(target_phone or cand_phone)
                if not match_name:
                    return {"msg": "update 模式需要提供候选人姓名"}
                names = df["候选人姓名"].map(_norm)
                mask = (names == match_name)
                if not target_name:
                    phones = df["手机号"].map(_norm)
                    if match_phone and match_phone != "未获取":
                        mask = mask | (phones == match_phone)
            if mask.sum() == 0:
                return {"msg": f"未找到候选人【{match_name}】的记录，无法更新"}
            def _merge(col, new_val):
                new_val = _norm(new_val)
                if not new_val or new_val == "未获取":
                    return
                df.loc[mask, col] = new_val
            _merge("候选人姓名", cand_name)
            _merge("手机号", cand_phone)
            _merge("期望薪资", parse_data.get("expect_salary", ""))
            _merge("到岗时间", parse_data.get("available_time", ""))
            _merge("岗位名称", post_name)
            _merge("招聘渠道", channel)
            old_note = _norm(df.loc[mask, "备注"].iloc[0])
            df.loc[mask, "备注"] = f"{old_note}\n【更新】{resume_text[:200]}"
            save_excel(df, user)
            return {"msg": f"已更新候选人【{match_name}】的信息", "updated": True, "parsed": parse_data}
        dup_mask = _dup_mask(df, cand_name, cand_phone)
        if mode == "auto" and dup_mask.sum() > 0:
            dup_rows = []
            for i in df.index[dup_mask]:
                dup_rows.append({
                    "name": _norm(df.loc[i, "候选人姓名"]),
                    "phone": _norm(df.loc[i, "手机号"]),
                    "expect_salary": _norm(df.loc[i, "期望薪资"]),
                    "available_time": _norm(df.loc[i, "到岗时间"]),
                    "status": _norm(df.loc[i, "沟通状态"]),
                })
            return {"msg": "检测到该候选人已存在，未写入台账", "duplicate": True, "exists": dup_rows}
        new_row = {
            "岗位名称": post_name,
            "候选人姓名": cand_name,
            "手机号": cand_phone,
            "招聘渠道": channel,
            "期望薪资": _norm(parse_data.get("expect_salary", "")),
            "到岗时间": _norm(parse_data.get("available_time", "")),
            "沟通状态": "待沟通",
            "面试时间": "",
            "面试形式": "",
            "面试结果": "",
            "备注": resume_text[:300]
        }
        df.loc[len(df)] = new_row
        save_excel(df, user)
        return {"msg": "简历解析存入台账成功", "parsed": parse_data}

    # 7. 更新候选人状态
    if name == "update_candidate_status":
        df = load_excel(user)
        for col in ["沟通状态", "面试时间", "面试形式", "面试结果", "备注"]:
            df[col] = df[col].astype(str).replace("nan", "")
        key = args.get("key")
        if key is not None:
            if key < 0 or key >= len(df):
                return {"msg": "该候选人记录不存在或已被删除"}
            mask = pd.Series(False, index=df.index)
            mask.iloc[key] = True
            cand_name = _norm(df.iloc[key]["候选人姓名"])
        else:
            name = str(args.get("name", ""))
            mask = df["候选人姓名"] == name
            if mask.sum() == 0:
                return {"msg": f"没有找到姓名为【{name}】的候选人记录"}
            cand_name = name
        for col, arg in [("沟通状态", "status"), ("面试时间", "interview_time"),
                         ("面试形式", "interview_type"), ("面试结果", "result")]:
            val = args.get(arg, "")
            if val:
                df.loc[mask, col] = str(val)
        note = str(args.get("note", ""))
        if note.strip():
            old_note = df.loc[mask, "备注"].iloc[0]
            df.loc[mask, "备注"] = f"{old_note}\n【更新备注】{note}"
        save_excel(df, user)
        return {"msg": f"状态更新成功：{cand_name}", "candidate_name": cand_name}

    # 8. 删除候选人
    if name == "delete_candidate":
        key = args.get("key")
        df = load_excel(user)
        if len(df) == 0 or key is None or key < 0 or key >= len(df):
            return {"msg": "该记录不存在或已被删除"}
        row = df.iloc[key]
        deleted = {
            "name": _norm(row["候选人姓名"]),
            "phone": _norm(row["手机号"]),
            "post": _norm(row["岗位名称"]),
        }
        df = df.drop(index=df.index[key]).reset_index(drop=True)
        save_excel(df, user)
        return {"msg": f"已删除候选人【{deleted['name']}】（{deleted['phone']}）的记录", "deleted": deleted}

    return {"error": f"未知工具：{name}"}


# ========== 接口：Agent 对话式指挥 ==========
@app.post("/agent_chat")
async def agent_chat(messages: list = Body(..., embed=True, description="对话历史，如 [{'role':'user','content':'...'}]"),
                     user: str = Depends(get_current_user)):
    """对话式指挥：LLM 自主调用内部工具完成招聘任务（数据按当前用户隔离）"""
    sys_prompt = (
        "你是「招聘AI助手」，帮HR管理候选人台账。工具：list_candidates(列候选人)、"
        "get_all_records(台账明细)、get_jobs(岗位)、get_stats(统计)、gen_talk(话术)、"
        "parse_resume_add(简历入库)、update_candidate_status(更新状态)、delete_candidate(删除)。"
        "按用户指令自主调用工具，删除前先确认，最后用简短中文汇报结果。"
    )
    msgs = [{"role": "system", "content": sys_prompt}]
    # 只保留最近 6 条用户/助手消息，避免历史过长拖慢响应
    recent = [m for m in messages if m.get("role") in ("user", "assistant") and m.get("content")][-6:]
    for m in recent:
        msgs.append({"role": m["role"], "content": m["content"]})

    tool_log = []
    for _ in range(12):
        try:
            resp = client.chat.completions.create(
                model=CURRENT_MODEL,
                messages=msgs,
                tools=AGENT_TOOLS,
                tool_choice="auto",
                max_tokens=512,
            )
        except Exception as e:
            # 模型不支持 tools 时降级为普通问答
            resp = client.chat.completions.create(model=CURRENT_MODEL, messages=msgs, max_tokens=512)
            return {"reply": resp.choices[0].message.content or "", "tool_calls": tool_log, "degraded": True}

        msg = resp.choices[0].message
        tool_calls = getattr(msg, "tool_calls", None)
        if not tool_calls:
            return {"reply": msg.content or "", "tool_calls": tool_log}

        # 记录 assistant 的 tool_calls，再逐个执行
        assistant_tool_msgs = []
        for tc in tool_calls:
            try:
                args = json.loads(tc.function.arguments or "{}")
            except json.JSONDecodeError:
                args = {}
            assistant_tool_msgs.append({
                "id": tc.id, "type": "function",
                "function": {"name": tc.function.name, "arguments": json.dumps(args, ensure_ascii=False)}
            })
        msgs.append({"role": "assistant", "content": msg.content or "", "tool_calls": assistant_tool_msgs})

        for tc in tool_calls:
            name = tc.function.name
            try:
                args = json.loads(tc.function.arguments or "{}")
            except json.JSONDecodeError:
                args = {}
            result = _agent_tool(name, args, user)
            tool_log.append({"name": name, "args": args, "result": result})
            # 工具结果截断，避免大结果拖慢下一轮
            content = json.dumps(result, ensure_ascii=False)
            if len(content) > 4000:
                content = content[:4000] + "...[截断]"
            msgs.append({
                "role": "tool",
                "tool_call_id": tc.id,
                "name": name,
                "content": content
            })
    return {"reply": "这一步涉及多个操作，为安全起见请拆成更具体的指令重试。", "tool_calls": tool_log}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app,host="127.0.0.1",port=8010)
